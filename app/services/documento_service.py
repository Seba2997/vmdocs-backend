import uuid
import io
import csv
import logging
import os
import shutil

from openpyxl import Workbook, load_workbook
import PyPDF2
import docx as python_docx

from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session
from supabase import create_client, Client

from app.config import SUPABASE_URL, SUPABASE_KEY, SUPABASE_BUCKET, TAMANO_MAXIMO_BYTES, TIPOS_MIME_PERMITIDOS

# Configuracion del backend de almacenamiento
_STORAGE_BACKEND = os.getenv("STORAGE_BACKEND", "supabase")  # "local" | "supabase"
_UPLOAD_DIR      = os.getenv("UPLOAD_DIR", "/uploads")
from app.models.documento_model import Documento
from app.models.caso_model import Caso
from app.models.caso_usuario_model import CasoUsuario

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# CLIENTE SUPABASE
# ─────────────────────────────────────────────

def _get_supabase() -> Client:
    """Instancia el cliente de Supabase Storage."""
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise HTTPException(
            status_code=503,
            detail="El módulo de almacenamiento no está configurado correctamente."
        )
    return create_client(SUPABASE_URL, SUPABASE_KEY)


# ─────────────────────────────────────────────
# VALIDACIONES
# ─────────────────────────────────────────────

def _validar_archivo(archivo: UploadFile, tamano: int) -> None:
    """
    Valida el tipo MIME, la extensión y el tamaño del archivo.
    Lanza HTTPException 400 si no cumple con los requisitos.
    """
    # Validar MIME type (comparar contra los valores del dict, no las claves)
    if archivo.content_type not in TIPOS_MIME_PERMITIDOS.values():
        raise HTTPException(
            status_code=400,
            detail=(
                f"Tipo de archivo no permitido: '{archivo.content_type}'. "
                f"Solo se aceptan: {', '.join(TIPOS_MIME_PERMITIDOS.values())}."
            )
        )

    # Validar extensión
    nombre = archivo.filename or ""
    extension = nombre.rsplit(".", 1)[-1].lower() if "." in nombre else ""
    extensiones_permitidas = {ext.lstrip(".") for ext in TIPOS_MIME_PERMITIDOS}
    if extension not in extensiones_permitidas:
        raise HTTPException(
            status_code=400,
            detail=f"Extensión de archivo no permitida: '.{extension}'."
        )

    # Validar tamaño
    if tamano > TAMANO_MAXIMO_BYTES:
        limite_mb = TAMANO_MAXIMO_BYTES / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"El archivo supera el tamaño máximo permitido de {limite_mb:.0f} MB."
        )


def _obtener_caso_activo_o_404(db: Session, caso_id: int) -> Caso:
    """Retorna el caso si existe y está activo, lanza 404 en caso contrario."""
    caso = db.query(Caso).filter(Caso.id == caso_id, Caso.activo == True).first()
    if not caso:
        raise HTTPException(status_code=404, detail="Caso no encontrado o inactivo.")
    return caso


def _obtener_documento_activo_o_404(db: Session, documento_id: int) -> Documento:
    """Retorna el documento si existe y está activo, lanza 404 en caso contrario."""
    doc = db.query(Documento).filter(Documento.id == documento_id, Documento.activo == True).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    return doc


# ─────────────────────────────────────────────
# CONTROL DE ACCESO
# ─────────────────────────────────────────────

def usuario_tiene_acceso_a_caso(db: Session, usuario_id: int, caso_id: int) -> bool:
    """Verifica si el usuario está asignado al caso."""
    return (
        db.query(CasoUsuario)
        .filter(
            CasoUsuario.caso_id == caso_id,
            CasoUsuario.usuario_id == usuario_id,
        )
        .first()
    ) is not None


def verificar_acceso_a_caso_o_403(db: Session, usuario_id: int, caso_id: int) -> None:
    """Lanza 403 si el usuario no tiene acceso al caso."""
    if not usuario_tiene_acceso_a_caso(db, usuario_id, caso_id):
        raise HTTPException(
            status_code=403,
            detail="No tienes permiso para acceder a este caso."
        )


# ─────────────────────────────────────────────
# EXTRACCIÓN DE TEXTO (para IA)
# ─────────────────────────────────────────────

def _extraer_texto_pdf(contenido: bytes) -> str | None:
    """
    Extrae texto plano de un PDF con texto embebido usando PyPDF2.
    Retorna None si el PDF es una imagen escaneada o si ocurre cualquier error.
    """
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(contenido))
        partes = [
            page.extract_text()
            for page in reader.pages
            if page.extract_text()
        ]
        texto_completo = "\n".join(partes).strip()
        return texto_completo[:20000] if texto_completo else None
    except Exception as exc:
        logger.warning("No se pudo extraer texto del PDF: %s", exc)
        return None


def _extraer_texto_xlsx(contenido: bytes) -> str | None:
    """
    Extrae texto plano de un XLSX (y de los CSV convertidos a XLSX)
    leyendo celda por celda con openpyxl.
    Solo considera hojas activas; celdas vacías se omiten.
    """
    try:
        wb = load_workbook(io.BytesIO(contenido), read_only=True, data_only=True)
        filas_texto = []
        for sheet in wb.worksheets:
            for fila in sheet.iter_rows(values_only=True):
                celda_str = "\t".join(
                    str(celda) for celda in fila if celda is not None
                )
                if celda_str.strip():
                    filas_texto.append(celda_str)
        wb.close()
        texto_completo = "\n".join(filas_texto).strip()
        return texto_completo[:20000] if texto_completo else None
    except Exception as exc:
        logger.warning("No se pudo extraer texto del XLSX: %s", exc)
        return None


def _extraer_texto_docx(contenido: bytes) -> str | None:
    """
    Extrae texto plano de un archivo DOCX usando python-docx.
    Recorre párrafos y tablas para cubrir el contenido completo del documento.
    """
    try:
        doc = python_docx.Document(io.BytesIO(contenido))
        partes = []

        # Párrafos normales
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                partes.append(parrafo.text.strip())

        # Texto dentro de tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                fila_texto = "\t".join(
                    celda.text.strip() for celda in fila.cells if celda.text.strip()
                )
                if fila_texto:
                    partes.append(fila_texto)

        texto_completo = "\n".join(partes).strip()
        return texto_completo[:20000] if texto_completo else None
    except Exception as exc:
        logger.warning("No se pudo extraer texto del DOCX: %s", exc)
        return None


def _extraer_texto_doc(contenido: bytes) -> str | None:
    """
    Extracción de texto de mejor esfuerzo para archivos DOC (formato Word 97-2003).
    El formato DOC es un binario OLE complejo; sin librerías externas como LibreOffice
    solo se puede recuperar texto visible de forma aproximada.
    Retorna None si no se encuentra texto legible.
    """
    try:
        # Decodificar bytes intentando latin-1 (cubre la mayoría de carácteres del formato OLE)
        raw = contenido.decode("latin-1", errors="ignore")
        # Filtrar solo caracteres imprimibles (letras, números, puntuación, espacios)
        texto = "".join(
            ch for ch in raw
            if ch.isprintable() or ch in ("\n", "\t", "\r")
        )
        # Eliminar líneas muy cortas (generalmente son artefactos binarios)
        lineas = [l.strip() for l in texto.splitlines() if len(l.strip()) > 3]
        texto_limpio = "\n".join(lineas).strip()
        return texto_limpio[:20000] if len(texto_limpio) > 50 else None
    except Exception as exc:
        logger.warning("No se pudo extraer texto del DOC: %s", exc)
        return None


def _extraer_texto(contenido: bytes, extension: str) -> str | None:
    """
    Dispatcher: selecciona la función de extracción correcta según la extensión.
    Retorna None para tipos sin soporte (imágenes, etc.).
    """
    if extension == "pdf":
        return _extraer_texto_pdf(contenido)
    if extension == "xlsx":
        return _extraer_texto_xlsx(contenido)
    if extension == "docx":
        return _extraer_texto_docx(contenido)
    if extension == "doc":
        return _extraer_texto_doc(contenido)
    return None  # png, jpg, jpeg → sin extracción


# ─────────────────────────────────────────────
# SUBIDA
# ─────────────────────────────────────────────

def subir_documento(
    db: Session,
    caso_id: int,
    usuario_id: int,
    archivo: UploadFile,
) -> Documento:
    """Sube archivo a Supabase y guarda metadata."""
    # 1. Validar que el caso existe y está activo
    _obtener_caso_activo_o_404(db, caso_id)

    # 2. Leer contenido para conocer el tamaño real
    contenido = archivo.file.read()
    tamano = len(contenido)

    # 3. Validar archivo (MIME, extensión, tamaño)
    _validar_archivo(archivo, tamano)

    # 4. Generar nombre único en storage y capturar info original
    nombre_original = archivo.filename or "archivo"
    extension = nombre_original.rsplit(".", 1)[-1].lower() if "." in nombre_original else "bin"
    content_type = archivo.content_type

    # TRANSFORMACIÓN CSV -> XLSX
    if extension == "csv":
        try:
            # Decodificar el CSV manejando posible BOM (muy común en archivos generados en Windows)
            text_content = contenido.decode('utf-8-sig')
            
            # Usar un método más robusto que el Sniffer para el delimitador (que a veces falla)
            # En Latinoamérica Excel exporta CSVs con ';' en lugar de ','.
            delimitador = ';' if text_content.count(';') > text_content.count(',') else ','
            
            csv_reader = csv.reader(io.StringIO(text_content), delimiter=delimitador)
            
            wb = Workbook()
            ws = wb.active
            for row in csv_reader:
                ws.append(row)
                
            out_stream = io.BytesIO()
            wb.save(out_stream)
            
            # Actualizar las variables para guardar el Excel generado en lugar del CSV
            contenido = out_stream.getvalue()
            tamano = len(contenido)
            extension = "xlsx"
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            
            # Modificar el nombre original para que la BD refleje que ahora es un Excel
            nombre_original = nombre_original.rsplit(".", 1)[0] + ".xlsx"
            
        except Exception as e:
            raise HTTPException(
                status_code=400, 
                detail=f"No se pudo procesar y transformar el archivo CSV a Excel: {str(e)}"
            )

    nombre_storage = f"casos/{caso_id}/{uuid.uuid4()}.{extension}"

    # 4b. Extraer texto para análisis IA según el tipo de archivo
    # PDF, XLSX (incl. CSV convertidos), DOCX, DOC → se extrae texto.
    # Imágenes → texto_extraido queda en None (la IA devolverá 422 si se intenta).
    texto_extraido: str | None = _extraer_texto(contenido, extension)

    # 5. Subir al backend de almacenamiento configurado
    if _STORAGE_BACKEND == "local":
        ruta_fisica = os.path.join(_UPLOAD_DIR, nombre_storage)
        os.makedirs(os.path.dirname(ruta_fisica), exist_ok=True)
        with open(ruta_fisica, "wb") as f:
            f.write(contenido)
    else:
        supabase = _get_supabase()
        try:
            supabase.storage.from_(SUPABASE_BUCKET).upload(
                path=nombre_storage,
                file=contenido,
                file_options={"content-type": content_type},
            )
        except Exception as exc:
            raise HTTPException(
                status_code=502,
                detail=f"Error al subir el archivo al almacenamiento: {exc}"
            )

    # 6. Persistir metadata en PostgreSQL
    nuevo_doc = Documento(
        nombre_original=nombre_original,
        nombre_storage=nombre_storage,
        tipo_mime=content_type,          # usar content_type local (ya actualizado si fue CSV→XLSX)
        tamano=tamano,
        caso_id=caso_id,
        usuario_id=usuario_id,
        texto_extraido=texto_extraido,  # puede ser None si el PDF no tiene texto
    )
    try:
        db.add(nuevo_doc)
        db.commit()
        db.refresh(nuevo_doc)
    except Exception:
        # Intentar eliminar el archivo de storage para no dejar huérfanos
        try:
            supabase.storage.from_(SUPABASE_BUCKET).remove([nombre_storage])
        except Exception:
            pass
        db.rollback()
        raise HTTPException(status_code=500, detail="Error al guardar la metadata del documento.")

    return nuevo_doc


# ─────────────────────────────────────────────
# LISTADO Y OBTENCIÓN
# ─────────────────────────────────────────────

def obtener_documentos_por_caso(db: Session, caso_id: int) -> list[Documento]:
    """Retorna los documentos activos de un caso activo."""
    _obtener_caso_activo_o_404(db, caso_id)
    return (
        db.query(Documento)
        .filter(Documento.caso_id == caso_id, Documento.activo == True)
        .order_by(Documento.fecha_subida.desc())
        .all()
    )


def obtener_documentos_inactivos_por_caso(db: Session, caso_id: int) -> list[Documento]:
    """Retorna los documentos inactivos (eliminados logicamente) de un caso activo."""
    _obtener_caso_activo_o_404(db, caso_id)
    return (
        db.query(Documento)
        .filter(Documento.caso_id == caso_id, Documento.activo == False)
        .order_by(Documento.fecha_subida.desc())
        .all()
    )


def obtener_documento_por_id(db: Session, documento_id: int) -> Documento:
    """Retorna un documento activo por su ID."""
    return _obtener_documento_activo_o_404(db, documento_id)


def obtener_documento_cualquier_estado(db: Session, documento_id: int) -> Documento:
    """Retorna un documento por ID sin importar si está activo o en papelera."""
    doc = db.query(Documento).filter(Documento.id == documento_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    return doc


# ─────────────────────────────────────────────
# SIGNED URL (descarga temporal)
# ─────────────────────────────────────────────

SIGNED_URL_EXPIRACION = 300  # segundos (5 minutos)


def generar_signed_url(db: Session, documento_id: int, modo: str = "descargar") -> dict:
    """
    Genera una URL para descargar o ver el documento.
    En modo local devuelve una URL directa al servidor.
    En produccion genera una URL firmada temporal de Supabase.
    """
    doc = _obtener_documento_activo_o_404(db, documento_id)

    if _STORAGE_BACKEND == "local":
        url = f"http://localhost:8000/uploads/{doc.nombre_storage}"
        return {"signed_url": url, "expira_en": 0}

    supabase = _get_supabase()
    opciones = {"download": doc.nombre_original} if modo == "descargar" else {}

    try:
        respuesta = supabase.storage.from_(SUPABASE_BUCKET).create_signed_url(
            path=doc.nombre_storage,
            expires_in=SIGNED_URL_EXPIRACION,
            options=opciones,
        )
        signed_url = respuesta.get("signedURL") or respuesta.get("signedUrl")
        if not signed_url:
            raise ValueError("La respuesta de Supabase no contiene la URL firmada.")
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Error al generar la URL de descarga: {exc}"
        )

    return {"signed_url": signed_url, "expira_en": SIGNED_URL_EXPIRACION}


# ─────────────────────────────────────────────
# SOFT DELETE / PAPELERA
# ─────────────────────────────────────────────

def toggle_estado_documento(
    db: Session,
    documento_id: int,
    usuario_solicitante_id: int | None = None,
) -> dict:
    """Activa o desactiva un documento (papelera)."""
    doc = db.query(Documento).filter(Documento.id == documento_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")

    # Verificar propiedad si es un usuario normal
    if usuario_solicitante_id is not None and doc.usuario_id != usuario_solicitante_id:
        raise HTTPException(
            status_code=403,
            detail="No tienes permiso para modificar este documento. Solo puedes gestionar los archivos que tú subiste.",
        )

    doc.activo = not doc.activo
    db.commit()
    db.refresh(doc)

    estado_texto = "restaurado" if doc.activo else "enviado a la papelera"
    return {"mensaje": f"Documento {estado_texto}.", "activo": doc.activo}


# ─────────────────────────────────────────────
# ELIMINACIÓN DEFINITIVA
# ─────────────────────────────────────────────

def eliminar_definitivamente(
    db: Session,
    documento_id: int,
    usuario_solicitante_id: int | None = None,
) -> dict:
    """Elimina permanentemente un documento de storage y BD."""
    # 1. Verificar existencia e inactividad
    doc = db.query(Documento).filter(Documento.id == documento_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")
    if doc.activo:
        raise HTTPException(
            status_code=409,
            detail="El documento está activo. Envíalo a la papelera antes de eliminarlo definitivamente.",
        )

    # 2. Verificar propiedad si es un usuario normal
    if usuario_solicitante_id is not None and doc.usuario_id != usuario_solicitante_id:
        raise HTTPException(
            status_code=403,
            detail="No tienes permiso para eliminar este documento. Solo puedes eliminar los archivos que tú subiste.",
        )

    nombre_storage = doc.nombre_storage

    # 3. Eliminar del backend de almacenamiento configurado
    if _STORAGE_BACKEND == "local":
        ruta_fisica = os.path.join(_UPLOAD_DIR, nombre_storage)
        try:
            if os.path.exists(ruta_fisica):
                os.remove(ruta_fisica)
        except Exception as exc:
            raise HTTPException(
                status_code=502,
                detail=f"Error al eliminar el archivo del disco: {exc}"
            )
    else:
        supabase = _get_supabase()
        try:
            supabase.storage.from_(SUPABASE_BUCKET).remove([nombre_storage])
        except Exception as exc:
            raise HTTPException(
                status_code=502,
                detail=f"Error al eliminar el archivo del almacenamiento: {exc}",
            )

    # 4. Eliminar el registro de PostgreSQL
    try:
        db.delete(doc)
        db.commit()
    except Exception:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail="El archivo se eliminó del almacenamiento pero no se pudo borrar el registro. Contacte al administrador.",
        )

    return {"mensaje": f"Documento '{doc.nombre_original}' eliminado definitivamente."}
